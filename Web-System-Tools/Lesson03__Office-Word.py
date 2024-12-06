from docx.api import Document
from pprint import pprint
from os.path import exists


def Read_DOCx():
    # Tải bảng đầu tiên từ tài liệu
    document = Document('data/demo.docx')
    table = document.tables[0]  # 0, 1, 2,... là số thứ tự của table trong file word

    # Dữ liệu sẽ là danh sách các hàng được biểu diễn dưới dạng từ điển
    # chứa dữ liệu của mỗi hàng.
    data = []

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        # Thiết lập ánh xạ dựa trên hàng đầu tiên
        # tiêu đề; những thứ này sẽ trở thành key của từ điển của chúng ta
        if i == 0:
            keys = tuple(text)
            continue

        # Xây dựng từ điển cho hàng này, ánh xạ key cho các giá trị cho hàng này
        row_data = dict(zip(keys, text))
        data.append(row_data)

    pprint(data)
    return data


from docx import Document
from docx.shared import Inches
import os


def Write_Docx():
    """
    Đây là đoạn chương trình tạo một file word, thêm các thông tin vào và lưu thành tệp
    Để chạy được bài này, cần có một thư mục `data`, trong đó có một tấm ảnh có tên image.png trong thư mục data: data/image.png
    """

    mDoc = Document()

    mDoc.add_heading('Tiêu đề', 0)

    p = mDoc.add_paragraph('Đây là một đoạn văn bản gồm có các kiểu chữ ')
    p.add_run('bold').bold = True
    p.add_run(' và ')
    p.add_run('italic.').italic = True

    mDoc.add_heading('Tiêu đề cấp 1 - Heading, level 1', level=1)
    mDoc.add_paragraph('Trích dẫn', style='Intense Quote')

    mDoc.add_paragraph('Gạch đầu dòng', style='List Bullet')
    mDoc.add_paragraph('Số thứ tự', style='List Number')

    os.makedirs('data', exist_ok=True)
    if exists('data/images.png'):
        mDoc.add_picture('data/image.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = mDoc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    mDoc.add_page_break()

    mDoc.save('data/demo.docx')
    print('Đã ghi một tệp vào thư mục data')


if __name__ == "__main__":
    Write_Docx()
    Read_DOCx()
